import os
import sys
import tempfile
import types
from pathlib import Path

import pytest

import chunker


def test_chunk_text_empty_or_whitespace_returns_empty():
    assert chunker.chunk_text("") == []
    assert chunker.chunk_text("   \n  \t") == []


def test_chunk_text_smaller_than_chunk_size_returns_single_chunk():
    text = "hello world"
    assert chunker.chunk_text(text, chunk_size=50, overlap=10) == [text]


def test_chunk_text_produces_overlapping_chunks():
    text = "a" * 900
    chunks = chunker.chunk_text(text, chunk_size=400, overlap=100)

    assert len(chunks) == 3
    assert chunks[0] == "a" * 400
    assert chunks[1] == "a" * 400
    assert chunks[2] == "a" * 100
    assert chunks[1][:100] == chunks[0][300:400]


def test_extract_text_reads_allowed_text_file(tmp_path):
    path = tmp_path / "sample.txt"
    content = "Line 1\nLine 2"
    path.write_text(content, encoding="utf-8")

    assert chunker.extract_text(str(path)) == content


def test_extract_text_returns_none_for_unsupported_extension(tmp_path):
    path = tmp_path / "secret.key"
    path.write_text("secret", encoding="utf-8")

    assert chunker.extract_text(str(path)) is None


def test_extract_text_returns_none_for_missing_file():
    assert chunker.extract_text("/this/path/does/not/exist.txt") is None


def test_extract_text_reads_docx_file_if_available(tmp_path):
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx is not installed")

    path = tmp_path / "document.docx"
    doc = docx.Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.save(path)

    assert chunker.extract_text(str(path)) == "alpha\nbeta"


def test_extract_text_reads_pdf_via_pypdf(monkeypatch, tmp_path):
    pdf_path = tmp_path / "document.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%EOF\n")

    class DummyPage:
        def extract_text(self):
            return "page text"

    class DummyReader:
        def __init__(self, filepath):
            self.pages = [DummyPage(), DummyPage()]

    stub_module = types.SimpleNamespace(PdfReader=DummyReader)
    monkeypatch.setitem(sys.modules, "pypdf", stub_module)

    assert chunker.extract_text(str(pdf_path)) == "page text\npage text"


def test_walk_and_chunk_yields_expected_items(tmp_path):
    root = tmp_path
    good_path = root / "good.txt"
    good_path.write_text("x" * 900, encoding="utf-8")
    skip_path = root / "skip.pem"
    skip_path.write_text("secret", encoding="utf-8")

    items = list(chunker.walk_and_chunk(str(root)))
    assert len(items) == 3

    for item in items:
        assert item["filepath"] == str(good_path)
        assert item["chunk_index"] in {0, 1, 2}
        assert item["text"]
        assert isinstance(item["mtime"], float)

    assert [item["chunk_index"] for item in items] == [0, 1, 2]
